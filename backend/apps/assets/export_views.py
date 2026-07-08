from django.http import HttpResponse
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated
from .models import Asset
from datetime import datetime
import csv
import openpyxl
from openpyxl.styles import Font, PatternFill
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def download_template(request):
    """Download a blank template for bulk asset import."""
    export_format = request.GET.get('format', 'xlsx').lower()
    
    # Template headers
    headers = [
        'name', 'asset_type', 'description', 'serial_number', 'model_number',
        'manufacturer', 'purchase_date', 'purchase_cost', 'current_value',
        'campus_code', 'building_code', 'floor_number', 'room_number',
        'status', 'condition', 'assigned_to_email', 'notes'
    ]
    
    # Sample data row
    sample_data = [
        'Dell Laptop XPS 15',
        'EQP',
        'High-performance laptop for engineering department',
        'SN123456789',
        'XPS-15-9520',
        'Dell',
        '2024-01-15',
        '1500.00',
        '1500.00',
        'MAIN',
        'ENG',
        '3',
        '301',
        'AVAILABLE',
        'EXCELLENT',
        'user@example.com',
        'New purchase for Q1 2024'
    ]
    
    if export_format == 'csv':
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="asset_import_template.csv"'
        
        writer = csv.writer(response)
        writer.writerow(headers)
        writer.writerow(sample_data)
        
        # Add instruction rows
        writer.writerow([])
        writer.writerow(['INSTRUCTIONS:'])
        writer.writerow(['1. Fill in asset data starting from row 2 (delete the sample row)'])
        writer.writerow(['2. Required fields: name, asset_type, campus_code, status'])
        writer.writerow(['3. asset_type values: EQP (Equipment), FUR (Furniture), VEH (Vehicle), BLD (Building), OTH (Other)'])
        writer.writerow(['4. status values: AVAILABLE, IN_USE, MAINTENANCE, RETIRED, DISPOSED'])
        writer.writerow(['5. condition values: EXCELLENT, GOOD, FAIR, POOR'])
        writer.writerow(['6. Date format: YYYY-MM-DD'])
        writer.writerow(['7. assigned_to_email must match an existing user email'])
        
        return response
        
    elif export_format == 'xlsx':
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = 'Asset Import Template'
        
        # Header styling
        header_font = Font(bold=True, color='FFFFFF', size=11)
        header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
        
        # Write headers
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
        
        # Write sample data
        for col, value in enumerate(sample_data, 1):
            ws.cell(row=2, column=col, value=value)
        
        # Auto-adjust column widths
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if cell.value and len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
        
        # Add instructions sheet
        ws_instructions = wb.create_sheet('Instructions')
        instructions = [
            ['Asset Import Template - Instructions'],
            [],
            ['REQUIRED FIELDS:'],
            ['• name - Asset name (text)'],
            ['• asset_type - Asset type code (see valid values below)'],
            ['• campus_code - Campus code where asset is located'],
            ['• status - Current status of the asset (see valid values below)'],
            [],
            ['VALID ASSET TYPES:'],
            ['• EQP - Equipment'],
            ['• FUR - Furniture'],
            ['• VEH - Vehicle'],
            ['• BLD - Building'],
            ['• OTH - Other'],
            [],
            ['VALID STATUS VALUES:'],
            ['• AVAILABLE - Asset is available for use'],
            ['• IN_USE - Asset is currently in use'],
            ['• MAINTENANCE - Asset is under maintenance'],
            ['• RETIRED - Asset is retired'],
            ['• DISPOSED - Asset has been disposed'],
            [],
            ['VALID CONDITION VALUES:'],
            ['• EXCELLENT - Like new condition'],
            ['• GOOD - Minor wear and tear'],
            ['• FAIR - Noticeable wear, fully functional'],
            ['• POOR - Significant wear, may need repair'],
            [],
            ['DATE FORMAT:'],
            ['• Use YYYY-MM-DD format (e.g., 2024-01-15)'],
            [],
            ['NOTES:'],
            ['• Delete the sample data row before importing'],
            ['• assigned_to_email must match an existing user in the system'],
            ['• campus_code, building_code must exist in the system'],
            ['• Numeric fields (purchase_cost, current_value) should be numbers without currency symbols'],
        ]
        
        for row_idx, instruction in enumerate(instructions, 1):
            ws_instructions.cell(row=row_idx, column=1, value=instruction[0])
            if row_idx == 1:
                ws_instructions.cell(row=row_idx, column=1).font = Font(bold=True, size=14)
        
        ws_instructions.column_dimensions['A'].width = 80
        
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = 'attachment; filename="asset_import_template.xlsx"'
        wb.save(response)
        
        return response
    
    else:
        return HttpResponse('Invalid format. Supported: csv, xlsx', status=400)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def export_assets(request):
    """Export assets in various formats (CSV, Excel, PDF, Word)."""
    from apps.core.utils import log_action, get_client_ip
    
    export_format = request.GET.get('format', 'csv').lower()
    
    # Get filtered assets
    assets = Asset.objects.select_related('campus', 'room', 'assigned_to').all()
    
    # Apply filters if provided
    status = request.GET.get('status')
    asset_type = request.GET.get('asset_type')
    
    if status:
        assets = assets.filter(status=status)
    if asset_type:
        assets = assets.filter(asset_type=asset_type)
    
    # BR-DM-03: Log data export with authorization
    log_action(
        user=request.user,
        action='EXPORT',
        model_name='Asset',
        object_id=None,
        details={
            'export_format': export_format,
            'record_count': assets.count(),
            'filters': {
                'status': status,
                'asset_type': asset_type
            },
            'authorized_by': request.user.username,
            'authorization_level': request.user.role
        },
        ip_address=get_client_ip(request)
    )
    
    if export_format == 'csv':
        # CSV Export
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = f'attachment; filename="assets-{datetime.now().strftime("%Y%m%d")}.csv"'
        
        writer = csv.writer(response)
        writer.writerow(['Asset ID', 'Name', 'Type', 'Status', 'Campus', 'Location', 'Purchase Date',
                       'Purchase Cost', 'Current Value', 'Assigned To'])
        
        for asset in assets:
            writer.writerow([
                asset.asset_id,
                asset.name,
                asset.get_asset_type_display(),
                asset.get_status_display(),
                asset.campus.name if asset.campus else '',
                str(asset.room) if asset.room else '',
                asset.purchase_date or '',
                asset.purchase_cost or '',
                asset.current_value or '',
                asset.assigned_to.get_full_name() if asset.assigned_to else ''
            ])
        
        return response
        
    elif export_format == 'excel':
        # Excel Export
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = 'Assets'
        
        # Header styling
        header_font = openpyxl.styles.Font(bold=True, color='FFFFFF')
        header_fill = openpyxl.styles.PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
        
        headers = ['Asset ID', 'Name', 'Type', 'Status', 'Campus', 'Location', 'Purchase Date',
                  'Purchase Cost', 'Current Value', 'Assigned To']
        
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
        
        # Data rows
        for row, asset in enumerate(assets, 2):
            ws.cell(row=row, column=1, value=asset.asset_id)
            ws.cell(row=row, column=2, value=asset.name)
            ws.cell(row=row, column=3, value=asset.get_asset_type_display())
            ws.cell(row=row, column=4, value=asset.get_status_display())
            ws.cell(row=row, column=5, value=asset.campus.name if asset.campus else '')
            ws.cell(row=row, column=6, value=str(asset.room) if asset.room else '')
            ws.cell(row=row, column=7, value=str(asset.purchase_date) if asset.purchase_date else '')
            ws.cell(row=row, column=8, value=float(asset.purchase_cost) if asset.purchase_cost else '')
            ws.cell(row=row, column=9, value=float(asset.current_value) if asset.current_value else '')
            ws.cell(row=row, column=10, value=asset.assigned_to.get_full_name() if asset.assigned_to else '')
        
        # Auto-adjust column widths
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
        
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = f'attachment; filename="assets-{datetime.now().strftime("%Y%m%d")}.xlsx"'
        wb.save(response)
        
        return response
        
    elif export_format == 'pdf':
        # PDF Export
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="assets-{datetime.now().strftime("%Y%m%d")}.pdf"'
        
        doc = SimpleDocTemplate(response, pagesize=A4)
        elements = []
        styles = getSampleStyleSheet()
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1e40af'),
            spaceAfter=30,
            alignment=1
        )
        elements.append(Paragraph('Asset Inventory Report', title_style))
        elements.append(Paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}', styles['Normal']))
        elements.append(Spacer(1, 0.3*inch))
        
        # Table data
        data = [['Asset ID', 'Name', 'Type', 'Status', 'Campus']]
        
        for asset in assets:
            data.append([
                asset.asset_id,
                asset.name[:30],
                asset.get_asset_type_display(),
                asset.get_status_display(),
                asset.campus.name if asset.campus else ''
            ])
        
        # Create table
        table = Table(data, repeatRows=1)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4472C4')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONTSIZE', (0, 1), (-1, -1), 8),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey])
        ]))
        
        elements.append(table)
        doc.build(elements)
        
        return response
        
    elif export_format == 'word':
        # Word Export
        doc = Document()
        
        # Title
        title = doc.add_heading('Asset Inventory Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
        doc.add_paragraph(f'Total Assets: {assets.count()}')
        doc.add_paragraph('')
        
        # Table
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = table.rows[0].cells
        headers = ['Asset ID', 'Name', 'Type', 'Status', 'Campus', 'Location']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)
        
        # Data rows
        for asset in assets:
            row_cells = table.add_row().cells
            row_cells[0].text = asset.asset_id
            row_cells[1].text = asset.name
            row_cells[2].text = asset.get_asset_type_display()
            row_cells[3].text = asset.get_status_display()
            row_cells[4].text = asset.campus.name if asset.campus else ''
            row_cells[5].text = str(asset.room) if asset.room else ''
        
        # Save to response
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="assets-{datetime.now().strftime("%Y%m%d")}.docx"'
        doc.save(response)
        
        return response
        
    else:
        return HttpResponse('Invalid format. Supported: csv, excel, pdf, word', status=400)
