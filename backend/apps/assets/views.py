from rest_framework import viewsets, status
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from django.db.models import Q
from apps.core.permissions import IsPropertyManager
from apps.core.utils import log_action, get_client_ip
from .models import (Campus, Building, Floor, Room, Asset, AssetTransfer,
                     AssetWarranty, AssetInsurance, AssetCheckout, AssetDocument,
                     Budget, BudgetTransaction, CheckoutExtensionRequest, AssetEvent)
from .serializers import (CampusSerializer, BuildingSerializer, FloorSerializer, 
                          RoomSerializer, AssetSerializer, AssetTransferSerializer,
                          AssetWarrantySerializer, AssetInsuranceSerializer, AssetCheckoutSerializer,
                          AssetDocumentSerializer, BudgetSerializer, BudgetTransactionSerializer,
                          CheckoutExtensionRequestSerializer)
import csv
import io
from django.db import transaction
from datetime import datetime
import openpyxl


class CampusViewSet(viewsets.ModelViewSet):
    queryset = Campus.objects.all()
    serializer_class = CampusSerializer
    permission_classes = [IsAuthenticated]


class BuildingViewSet(viewsets.ModelViewSet):
    queryset = Building.objects.all()
    serializer_class = BuildingSerializer
    permission_classes = [IsAuthenticated]
    filterset_fields = ['campus']


class FloorViewSet(viewsets.ModelViewSet):
    queryset = Floor.objects.all()
    serializer_class = FloorSerializer
    permission_classes = [IsAuthenticated]
    filterset_fields = ['building']


class RoomViewSet(viewsets.ModelViewSet):
    queryset = Room.objects.all()
    serializer_class = RoomSerializer
    permission_classes = [IsAuthenticated]
    filterset_fields = ['floor', 'room_type']
    search_fields = ['number', 'name']


class AssetViewSet(viewsets.ModelViewSet):
    queryset = Asset.objects.select_related('campus', 'room', 'assigned_to').all()
    serializer_class = AssetSerializer
    permission_classes = [IsAuthenticated]
    filterset_fields = ['status', 'asset_type', 'campus', 'manufacturer', 'condition']
    search_fields = ['asset_id', 'name', 'description', 'serial_number', 'model_number', 'manufacturer']
    
    def perform_create(self, serializer):
        asset = serializer.save()
        asset._request_user = self.request.user
        asset.save()
    
    def perform_update(self, serializer):
        asset = serializer.save()
        asset._request_user = self.request.user
        asset.save()
    
    @action(detail=True, methods=['post'], permission_classes=[IsPropertyManager])
    def transfer(self, request, pk=None):
        """Transfer asset to a new location."""
        asset = self.get_object()
        to_room_id = request.data.get('to_room')
        reason = request.data.get('reason', '')
        
        if not to_room_id:
            return Response({'error': 'to_room is required'}, status=status.HTTP_400_BAD_REQUEST)
        
        transfer = AssetTransfer.objects.create(
            asset=asset,
            from_room=asset.room,
            to_room_id=to_room_id,
            reason=reason,
            requested_by=request.user,
            approved_by=request.user
        )
        transfer._request_user = request.user
        
        asset.room_id = to_room_id
        asset.save()
        
        return Response(AssetTransferSerializer(transfer).data)
    
    @action(detail=False, methods=['post'], permission_classes=[IsPropertyManager])
    def bulk_status_update(self, request):
        """Bulk update asset statuses."""
        asset_ids = request.data.get('asset_ids', [])
        new_status = request.data.get('status')
        
        if not asset_ids or not new_status:
            return Response(
                {'error': 'asset_ids and status are required'}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        if new_status not in dict(Asset.STATUS_CHOICES):
            return Response(
                {'error': f'Invalid status. Must be one of: {", ".join(dict(Asset.STATUS_CHOICES).keys())}'}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        try:
            with transaction.atomic():
                # Store original states for rollback capability
                original_states = []
                assets = Asset.objects.filter(asset_id__in=asset_ids)
                
                for asset in assets:
                    original_states.append({
                        'asset_id': asset.asset_id,
                        'old_status': asset.status
                    })
                    asset.status = new_status
                
                Asset.objects.bulk_update(assets, ['status'])
                
                # Log the bulk operation
                log_action(
                    user=request.user,
                    action='BULK_UPDATE',
                    model_name='Asset',
                    object_id=None,
                    details=f'Updated {len(assets)} assets to status {new_status}',
                    ip_address=get_client_ip(request)
                )
                
                return Response({
                    'success': True,
                    'updated_count': len(assets),
                    'original_states': original_states,
                    'message': f'Successfully updated {len(assets)} assets to {new_status}'
                })
        except Exception as e:
            return Response(
                {'error': str(e)}, 
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
    
    @action(detail=False, methods=['get'], permission_classes=[IsPropertyManager], url_path='download-template')
    def download_import_template(self, request):
        """Download CSV or Excel template for bulk import."""
        format_type = request.query_params.get('format', 'csv').lower()
        
        # Define template headers with all fields
        headers = [
            'name',
            'asset_type',
            'campus_code',
            'status',
            'description',
            'manufacturer',
            'model_number',
            'serial_number',
            'condition',
            'supplier',
            'purchase_date',
            'purchase_cost',
            'current_value',
            'room_code',
        ]
        
        # Sample data row
        sample_data = [
            'Dell Laptop XPS 15',
            'EQP',
            'BURIE',
            'AVAILABLE',
            'High-performance laptop for office use',
            'Dell',
            'XPS 15 9520',
            'SN123456789',
            'EXCELLENT',
            'Tech Solutions Ltd',
            '2024-01-15',
            '45000',
            '45000',
            'B1-F2-R101',
        ]
        
        # Instructions row
        instructions = [
            'Required: Asset name',
            'Required: EQP, FUR, VEH, BLD, OTH',
            'Required: Campus code (e.g., BURIE)',
            'Required: AVAILABLE, IN_USE, UNDER_MAINTENANCE, CONDEMNED',
            'Optional: Detailed description',
            'Optional: Manufacturer or brand',
            'Optional: Model number',
            'Optional: Unique serial number',
            'Optional: EXCELLENT, GOOD, FAIR, POOR',
            'Optional: Supplier/vendor name',
            'Optional: Format YYYY-MM-DD',
            'Optional: Purchase cost in ETB',
            'Optional: Current value in ETB',
            'Optional: Room code (Building-Floor-Room)',
        ]
        
        if format_type == 'csv':
            # Generate CSV
            from django.http import HttpResponse
            response = HttpResponse(content_type='text/csv')
            response['Content-Disposition'] = 'attachment; filename="asset_import_template.csv"'
            
            writer = csv.writer(response)
            writer.writerow(headers)
            writer.writerow(instructions)
            writer.writerow(sample_data)
            
            return response
            
        elif format_type in ['xlsx', 'excel']:
            # Generate Excel
            from django.http import HttpResponse
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
            
            wb = Workbook()
            ws = wb.active
            ws.title = "Asset Import Template"
            
            # Header row styling
            header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
            header_font = Font(bold=True, color="FFFFFF", size=11)
            instruction_fill = PatternFill(start_color="E7E6E6", end_color="E7E6E6", fill_type="solid")
            instruction_font = Font(italic=True, size=9)
            border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            
            # Write headers
            for col_idx, header in enumerate(headers, start=1):
                cell = ws.cell(row=1, column=col_idx, value=header)
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal='center', vertical='center')
                cell.border = border
                ws.column_dimensions[cell.column_letter].width = 20
            
            # Write instructions
            for col_idx, instruction in enumerate(instructions, start=1):
                cell = ws.cell(row=2, column=col_idx, value=instruction)
                cell.fill = instruction_fill
                cell.font = instruction_font
                cell.alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
                cell.border = border
            
            # Write sample data
            for col_idx, data in enumerate(sample_data, start=1):
                cell = ws.cell(row=3, column=col_idx, value=data)
                cell.border = border
                cell.alignment = Alignment(horizontal='left', vertical='center')
            
            # Adjust row heights
            ws.row_dimensions[1].height = 25
            ws.row_dimensions[2].height = 40
            ws.row_dimensions[3].height = 20
            
            # Freeze header row
            ws.freeze_panes = 'A3'
            
            # Create response
            from io import BytesIO
            buffer = BytesIO()
            wb.save(buffer)
            buffer.seek(0)
            
            response = HttpResponse(
                buffer.read(),
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = 'attachment; filename="asset_import_template.xlsx"'
            
            return response
        
        else:
            return Response(
                {'error': 'Invalid format. Use csv or xlsx'}, 
                status=status.HTTP_400_BAD_REQUEST
            )
    
    @action(detail=False, methods=['post'], permission_classes=[IsPropertyManager])
    def bulk_import(self, request):
        """Bulk import assets from CSV or Excel with enhanced field support."""
        file = request.FILES.get('file')
        
        if not file:
            return Response(
                {'error': 'File is required'}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        file_extension = file.name.split('.')[-1].lower()
        
        if file_extension not in ['csv', 'xls', 'xlsx']:
            return Response(
                {'error': 'File must be CSV or Excel (.xls, .xlsx)'}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        try:
            # Parse file based on type
            rows = []
            
            if file_extension == 'csv':
                # Read CSV
                decoded_file = file.read().decode('utf-8')
                io_string = io.StringIO(decoded_file)
                reader = csv.DictReader(io_string)
                rows = list(reader)
                headers = reader.fieldnames
            else:
                # Read Excel
                workbook = openpyxl.load_workbook(file, read_only=True)
                sheet = workbook.active
                
                # Get headers from first row
                headers = [cell.value for cell in sheet[1]]
                
                # Get data rows (skip instruction row if present)
                start_row = 3 if sheet.max_row > 2 else 2
                for row in sheet.iter_rows(min_row=start_row, values_only=True):
                    row_dict = {}
                    for idx, value in enumerate(row):
                        if idx < len(headers) and headers[idx]:
                            row_dict[headers[idx]] = str(value) if value is not None else ''
                    if any(row_dict.values()):  # Skip empty rows
                        rows.append(row_dict)
            
            # Validate headers
            required_fields = ['name', 'asset_type', 'campus_code', 'status']
            
            missing_fields = [field for field in required_fields if field not in headers]
            if missing_fields:
                return Response(
                    {'error': f'Missing required fields: {", ".join(missing_fields)}'}, 
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            # Preview mode - just validate and return preview
            preview_mode = request.data.get('preview', 'false').lower() == 'true'
            
            preview_data = []
            errors = []
            
            for idx, row in enumerate(rows, start=1):
                row_data = {
                    'row': idx,
                    'name': row.get('name', '').strip(),
                    'asset_type': row.get('asset_type', '').strip().upper(),
                    'campus_code': row.get('campus_code', '').strip().upper(),
                    'status': row.get('status', '').strip().upper(),
                    'description': row.get('description', '').strip(),
                    'manufacturer': row.get('manufacturer', '').strip(),
                    'model_number': row.get('model_number', '').strip(),
                    'serial_number': row.get('serial_number', '').strip(),
                    'condition': row.get('condition', '').strip().upper() or 'GOOD',
                    'supplier': row.get('supplier', '').strip(),
                    'purchase_date': row.get('purchase_date', '').strip(),
                    'purchase_cost': row.get('purchase_cost', '').strip(),
                    'current_value': row.get('current_value', '').strip(),
                    'room_code': row.get('room_code', '').strip(),
                }
                
                # Validate row
                row_errors = []
                
                if not row_data['name']:
                    row_errors.append('Name is required')
                
                if row_data['asset_type'] not in dict(Asset.TYPE_CHOICES):
                    row_errors.append(f'Invalid asset_type: {row_data["asset_type"]}. Must be one of: EQP, FUR, VEH, BLD, OTH')
                
                if row_data['status'] not in dict(Asset.STATUS_CHOICES):
                    row_errors.append(f'Invalid status: {row_data["status"]}. Must be one of: AVAILABLE, IN_USE, UNDER_MAINTENANCE, CONDEMNED')
                
                if row_data['condition'] and row_data['condition'] not in ['EXCELLENT', 'GOOD', 'FAIR', 'POOR']:
                    row_errors.append(f'Invalid condition: {row_data["condition"]}. Must be one of: EXCELLENT, GOOD, FAIR, POOR')
                
                # Validate campus
                try:
                    campus = Campus.objects.get(code=row_data['campus_code'])
                    row_data['campus_name'] = campus.name
                    row_data['campus_id'] = campus.id
                except Campus.DoesNotExist:
                    row_errors.append(f'Campus not found: {row_data["campus_code"]}')
                
                # Validate serial number uniqueness
                if row_data['serial_number']:
                    if Asset.objects.filter(serial_number=row_data['serial_number']).exists():
                        row_errors.append(f'Serial number already exists: {row_data["serial_number"]}')
                else:
                    # Validate similar fields for duplicate detection if no serial number is provided
                    if 'campus_id' in row_data:
                        duplicate_query = Asset.objects.filter(
                            name__iexact=row_data['name'],
                            asset_type=row_data['asset_type'],
                            campus_id=row_data['campus_id']
                        )
                        if row_data.get('manufacturer'):
                            duplicate_query = duplicate_query.filter(manufacturer__iexact=row_data['manufacturer'])
                        if row_data.get('model_number'):
                            duplicate_query = duplicate_query.filter(model_number__iexact=row_data['model_number'])
                            
                        if duplicate_query.exists():
                            row_errors.append('A highly similar asset (same name, type, campus, and manufacturer) already exists. Please verify to avoid duplication.')
                
                # Validate purchase date format
                if row_data['purchase_date']:
                    try:
                        from datetime import datetime
                        datetime.strptime(row_data['purchase_date'], '%Y-%m-%d')
                    except ValueError:
                        row_errors.append(f'Invalid date format: {row_data["purchase_date"]}. Use YYYY-MM-DD')
                
                # Validate numeric fields
                if row_data['purchase_cost']:
                    try:
                        float(row_data['purchase_cost'])
                    except ValueError:
                        row_errors.append(f'Invalid purchase_cost: {row_data["purchase_cost"]}')
                
                if row_data['current_value']:
                    try:
                        float(row_data['current_value'])
                    except ValueError:
                        row_errors.append(f'Invalid current_value: {row_data["current_value"]}')
                
                if row_errors:
                    errors.append({'row': idx, 'errors': row_errors, 'data': row_data})
                
                row_data['valid'] = len(row_errors) == 0
                preview_data.append(row_data)
            
            if preview_mode:
                return Response({
                    'preview': preview_data,
                    'total_rows': len(rows),
                    'valid_rows': len([r for r in preview_data if r['valid']]),
                    'invalid_rows': len([r for r in preview_data if not r['valid']]),
                    'errors': errors
                })
            
            # Actual import
            if errors:
                return Response(
                    {'error': 'File contains errors', 'errors': errors, 'total_errors': len(errors)}, 
                    status=status.HTTP_400_BAD_REQUEST
                )
            
            created_assets = []
            
            with transaction.atomic():
                for row_data in preview_data:
                    if not row_data['valid']:
                        continue
                    
                    campus = Campus.objects.get(code=row_data['campus_code'])
                    
                    # Prepare asset data
                    asset_data = {
                        'name': row_data['name'],
                        'asset_type': row_data['asset_type'],
                        'campus': campus,
                        'status': row_data['status'],
                        'condition': row_data['condition'],
                    }
                    
                    # Add optional fields
                    if row_data['description']:
                        asset_data['description'] = row_data['description']
                    if row_data['manufacturer']:
                        asset_data['manufacturer'] = row_data['manufacturer']
                    if row_data['model_number']:
                        asset_data['model_number'] = row_data['model_number']
                    if row_data['serial_number']:
                        asset_data['serial_number'] = row_data['serial_number']
                    if row_data['supplier']:
                        asset_data['supplier'] = row_data['supplier']
                    if row_data['purchase_date']:
                        asset_data['purchase_date'] = row_data['purchase_date']
                    if row_data['purchase_cost']:
                        asset_data['purchase_cost'] = float(row_data['purchase_cost'])
                    if row_data['current_value']:
                        asset_data['current_value'] = float(row_data['current_value'])
                    
                    asset = Asset.objects.create(**asset_data)
                    created_assets.append({
                        'asset_id': asset.asset_id,
                        'name': asset.name,
                        'row': row_data['row']
                    })
                
                # Log the bulk import
                log_action(
                    user=request.user,
                    action='BULK_IMPORT',
                    model_name='Asset',
                    object_id=None,
                    details=f'Imported {len(created_assets)} assets from {file_extension.upper()}',
                    ip_address=get_client_ip(request)
                )
            
            return Response({
                'success': True,
                'created': len(created_assets),
                'assets': created_assets,
                'message': f'Successfully imported {len(created_assets)} assets'
            })
            
        except Exception as e:
            import traceback
            return Response(
                {'error': f'Import failed: {str(e)}', 'traceback': traceback.format_exc()}, 
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
    
    @action(detail=False, methods=['post'], permission_classes=[IsPropertyManager])
    def rollback_bulk_update(self, request):
        """Rollback a bulk status update."""
        original_states = request.data.get('original_states', [])
        
        if not original_states:
            return Response(
                {'error': 'original_states is required'}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        try:
            with transaction.atomic():
                rollback_count = 0
                
                for state in original_states:
                    asset_id = state.get('asset_id')
                    old_status = state.get('old_status')
                    
                    if asset_id and old_status:
                        Asset.objects.filter(asset_id=asset_id).update(status=old_status)
                        rollback_count += 1
                
                # Log the rollback
                log_action(
                    user=request.user,
                    action='ROLLBACK',
                    model_name='Asset',
                    object_id=None,
                    details=f'Rolled back {rollback_count} asset status changes',
                    ip_address=get_client_ip(request)
                )
                
                return Response({
                    'success': True,
                    'rollback_count': rollback_count,
                    'message': f'Successfully rolled back {rollback_count} assets'
                })
        except Exception as e:
            return Response(
                {'error': str(e)}, 
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
    
    @action(detail=False, methods=['get'], permission_classes=[IsAuthenticated], url_path='export')
    def export(self, request):
        """Export assets in various formats (CSV, Excel, PDF, Word)."""
        from django.http import HttpResponse
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib import colors
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Support both DRF and Django request objects
        export_format = (getattr(request, 'query_params', None) or request.GET).get('format', 'csv').lower()

        # Get filtered assets
        queryset = self.filter_queryset(self.get_queryset())
        assets = queryset.select_related('campus', 'room', 'assigned_to')

        if export_format == 'csv':
            # CSV Export
            response = HttpResponse(content_type='text/csv')
            response['Content-Disposition'] = f'attachment; filename="assets-{datetime.now().strftime("%Y%m%d")}.csv"'

            writer = csv.writer(response)
            writer.writerow(['Asset ID', 'Name', 'Type', 'Status', 'Campus', 'Location', 'Purchase Date',
                           'Purchase Cost', 'Current Value', 'Assigned To'])

            for asset in assets:
                writer.writerow([
                    asset.asset_id,
                    asset.name,
                    asset.get_asset_type_display(),
                    asset.get_status_display(),
                    asset.campus.name if asset.campus else '',
                    str(asset.room) if asset.room else '',
                    asset.purchase_date or '',
                    asset.purchase_cost or '',
                    asset.current_value or '',
                    asset.assigned_to.get_full_name() if asset.assigned_to else ''
                ])

            return response

        elif export_format == 'excel':
            # Excel Export
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = 'Assets'

            # Header styling
            header_font = openpyxl.styles.Font(bold=True, color='FFFFFF')
            header_fill = openpyxl.styles.PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')

            headers = ['Asset ID', 'Name', 'Type', 'Status', 'Campus', 'Location', 'Purchase Date',
                      'Purchase Cost', 'Current Value', 'Assigned To']

            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col, value=header)
                cell.font = header_font
                cell.fill = header_fill

            # Data rows
            for row, asset in enumerate(assets, 2):
                ws.cell(row=row, column=1, value=asset.asset_id)
                ws.cell(row=row, column=2, value=asset.name)
                ws.cell(row=row, column=3, value=asset.get_asset_type_display())
                ws.cell(row=row, column=4, value=asset.get_status_display())
                ws.cell(row=row, column=5, value=asset.campus.name if asset.campus else '')
                ws.cell(row=row, column=6, value=str(asset.room) if asset.room else '')
                ws.cell(row=row, column=7, value=str(asset.purchase_date) if asset.purchase_date else '')
                ws.cell(row=row, column=8, value=float(asset.purchase_cost) if asset.purchase_cost else '')
                ws.cell(row=row, column=9, value=float(asset.current_value) if asset.current_value else '')
                ws.cell(row=row, column=10, value=asset.assigned_to.get_full_name() if asset.assigned_to else '')

            # Auto-adjust column widths
            for column in ws.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                ws.column_dimensions[column_letter].width = adjusted_width

            response = HttpResponse(
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = f'attachment; filename="assets-{datetime.now().strftime("%Y%m%d")}.xlsx"'
            wb.save(response)

            return response

        elif export_format == 'pdf':
            # PDF Export
            response = HttpResponse(content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="assets-{datetime.now().strftime("%Y%m%d")}.pdf"'

            doc = SimpleDocTemplate(response, pagesize=A4)
            elements = []
            styles = getSampleStyleSheet()

            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor=colors.HexColor('#1e40af'),
                spaceAfter=30,
                alignment=1
            )
            elements.append(Paragraph('Asset Inventory Report', title_style))
            elements.append(Paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}', styles['Normal']))
            elements.append(Spacer(1, 0.3*inch))

            # Table data
            data = [['Asset ID', 'Name', 'Type', 'Status', 'Campus']]

            for asset in assets:
                data.append([
                    asset.asset_id,
                    asset.name[:30],
                    asset.get_asset_type_display(),
                    asset.get_status_display(),
                    asset.campus.name if asset.campus else ''
                ])

            # Create table
            table = Table(data, repeatRows=1)
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4472C4')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey])
            ]))

            elements.append(table)
            doc.build(elements)

            return response

        elif export_format == 'word':
            # Word Export
            doc = Document()

            # Title
            title = doc.add_heading('Asset Inventory Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Metadata
            doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
            doc.add_paragraph(f'Total Assets: {assets.count()}')
            doc.add_paragraph('')

            # Table
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Light Grid Accent 1'

            # Header row
            header_cells = table.rows[0].cells
            headers = ['Asset ID', 'Name', 'Type', 'Status', 'Campus', 'Location']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(11)

            # Data rows
            for asset in assets:
                row_cells = table.add_row().cells
                row_cells[0].text = asset.asset_id
                row_cells[1].text = asset.name
                row_cells[2].text = asset.get_asset_type_display()
                row_cells[3].text = asset.get_status_display()
                row_cells[4].text = asset.campus.name if asset.campus else ''
                row_cells[5].text = str(asset.room) if asset.room else ''

            # Save to response
            response = HttpResponse(
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="assets-{datetime.now().strftime("%Y%m%d")}.docx"'
            doc.save(response)

            return response

        else:
            return Response(
                {'error': 'Invalid format. Supported formats: csv, excel, pdf, word'},
                status=status.HTTP_400_BAD_REQUEST
            )



class AssetWarrantyViewSet(viewsets.ModelViewSet):
    queryset = AssetWarranty.objects.select_related('asset').all()
    serializer_class = AssetWarrantySerializer
    permission_classes = [IsAuthenticated]
    filterset_fields = ['asset']
    
    @action(detail=False, methods=['get'])
    def expiring_soon(self, request):
        """Get warranties expiring in next 30 days."""
        from datetime import date, timedelta
        thirty_days = date.today() + timedelta(days=30)
        warranties = self.queryset.filter(end_date__lte=thirty_days, end_date__gte=date.today())
        serializer = self.get_serializer(warranties, many=True)
        return Response(serializer.data)


class AssetInsuranceViewSet(viewsets.ModelViewSet):
    queryset = AssetInsurance.objects.select_related('asset').all()
    serializer_class = AssetInsuranceSerializer
    permission_classes = [IsAuthenticated]
    filterset_fields = ['asset', 'policy_type']
    
    @action(detail=False, methods=['get'])
    def renewal_due(self, request):
        """Get insurance policies due for renewal in next 30 days."""
        from datetime import date, timedelta
        thirty_days = date.today() + timedelta(days=30)
        policies = self.queryset.filter(renewal_date__lte=thirty_days, renewal_date__gte=date.today())
        serializer = self.get_serializer(policies, many=True)
        return Response(serializer.data)


class AssetCheckoutViewSet(viewsets.ModelViewSet):
    queryset = AssetCheckout.objects.select_related('asset', 'checked_out_to', 'checked_out_by').all()
    serializer_class = AssetCheckoutSerializer
    permission_classes = [IsAuthenticated]
    filterset_fields = ['asset', 'checked_out_to', 'is_returned']
    
    def perform_create(self, serializer):
        """Create checkout and update asset status."""
        checkout = serializer.save(checked_out_by=self.request.user)
        
        # Update asset status to IN_USE and assign to user
        asset = checkout.asset
        asset.status = 'IN_USE'
        asset.assigned_to = checkout.checked_out_to
        asset.save()
        
        return checkout
    
    @action(detail=True, methods=['post'], permission_classes=[IsPropertyManager])
    def checkin(self, request, pk=None):
        """Check in an asset."""
        checkout = self.get_object()
        
        if checkout.is_returned:
            return Response({'error': 'Asset already checked in'}, status=status.HTTP_400_BAD_REQUEST)
        
        from django.utils import timezone
        
        # Wrap in transaction to ensure atomicity
        with transaction.atomic():
            checkout.actual_return_date = timezone.now()
            checkout.is_returned = True
            checkout.return_condition = request.data.get('return_condition', 'GOOD')
            checkout.notes = request.data.get('notes', checkout.notes)
            checkout.save()
            
            # Update asset status
            asset = checkout.asset
            asset.status = 'AVAILABLE'
            asset.assigned_to = None
            asset.save()
            
            # Create asset event
            AssetEvent.objects.create(
                asset=asset,
                event_type='CHECKED_IN',
                description=f'Asset checked in by {request.user.get_full_name()}. Return condition: {checkout.return_condition}',
                actor=request.user,
                related_checkout=checkout,
                event_data={
                    'return_condition': checkout.return_condition,
                    'notes': checkout.notes,
                    'checkout_id': checkout.id
                }
            )
            
            # Notify the user who checked out the asset
            from apps.core.models import Notification
            Notification.objects.create(
                user=checkout.checked_out_to,
                title='Asset Return Completed',
                message=f'Your return of {asset.asset_id} has been completed by {request.user.get_full_name()}',
                notification_type='SUCCESS'
            )
        
        return Response(AssetCheckoutSerializer(checkout).data)
    
    @action(detail=False, methods=['get'], permission_classes=[IsPropertyManager])
    def pending_returns(self, request):
        """Get checkouts with pending return requests (initiated by owners)."""
        from datetime import date, timedelta
        from django.utils import timezone
        
        # Get checkouts that are not returned and either:
        # 1. Have a return initiated event in the last 30 days
        # 2. Are overdue
        pending_checkouts = self.queryset.filter(
            is_returned=False
        ).select_related('asset', 'checked_out_to', 'checked_out_by')
        
        # Check for return initiated events
        from apps.assets.models import AssetEvent
        return_initiated_asset_ids = AssetEvent.objects.filter(
            event_type='RETURN_INITIATED',
            event_date__gte=timezone.now() - timedelta(days=30)
        ).values_list('asset_id', flat=True)
        
        # Filter to checkouts with return initiated or overdue
        pending_checkouts = pending_checkouts.filter(
            Q(asset_id__in=return_initiated_asset_ids) | 
            Q(expected_return_date__lt=date.today())
        ).order_by('expected_return_date')
        
        serializer = self.get_serializer(pending_checkouts, many=True)
        return Response(serializer.data)
    
    @action(detail=False, methods=['get'])
    def overdue(self, request):
        """Get overdue checkouts."""
        from datetime import date
        overdue_checkouts = self.queryset.filter(
            is_returned=False,
            expected_return_date__lt=date.today()
        )
        serializer = self.get_serializer(overdue_checkouts, many=True)
        return Response(serializer.data)


class AssetDocumentViewSet(viewsets.ModelViewSet):
    queryset = AssetDocument.objects.select_related('asset', 'uploaded_by').all()
    serializer_class = AssetDocumentSerializer
    permission_classes = [IsAuthenticated]
    filterset_fields = ['asset', 'document_type']
    
    def perform_create(self, serializer):
        serializer.save(uploaded_by=self.request.user)


class BudgetViewSet(viewsets.ModelViewSet):
    queryset = Budget.objects.select_related('campus').all()
    serializer_class = BudgetSerializer
    permission_classes = [IsPropertyManager]
    filterset_fields = ['fiscal_year', 'campus', 'is_active']
    
    @action(detail=True, methods=['post'])
    def add_transaction(self, request, pk=None):
        """Add a transaction to budget."""
        budget = self.get_object()
        
        transaction_data = {
            'budget': budget.id,
            'transaction_type': request.data.get('transaction_type'),
            'amount': request.data.get('amount'),
            'description': request.data.get('description'),
            'reference_number': request.data.get('reference_number', ''),
            'asset': request.data.get('asset'),
            'approved_by': request.user.id
        }
        
        serializer = BudgetTransactionSerializer(data=transaction_data)
        if serializer.is_valid():
            transaction = serializer.save()
            
            # Update budget spent amount
            budget.spent_amount += float(request.data.get('amount', 0))
            budget.save()
            
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    
    @action(detail=False, methods=['get'])
    def summary(self, request):
        """Get budget summary."""
        budgets = self.queryset.filter(is_active=True)
        total_budget = sum(b.total_amount for b in budgets)
        total_spent = sum(b.spent_amount for b in budgets)
        total_remaining = total_budget - total_spent
        
        return Response({
            'total_budget': total_budget,
            'total_spent': total_spent,
            'total_remaining': total_remaining,
            'utilization_percentage': (total_spent / total_budget * 100) if total_budget > 0 else 0,
            'active_budgets': budgets.count()
        })


class BudgetTransactionViewSet(viewsets.ReadOnlyModelViewSet):
    queryset = BudgetTransaction.objects.select_related('budget', 'asset', 'approved_by').all()
    serializer_class = BudgetTransactionSerializer
    permission_classes = [IsAuthenticated]
    filterset_fields = ['budget', 'transaction_type', 'asset']


class CheckoutExtensionRequestViewSet(viewsets.ModelViewSet):
    """Manage checkout extension requests with approval workflow."""
    queryset = CheckoutExtensionRequest.objects.select_related(
        'checkout__asset', 'requested_by', 'reviewed_by', 'checkout__checked_out_to'
    ).all()
    serializer_class = CheckoutExtensionRequestSerializer
    permission_classes = [IsAuthenticated]
    filterset_fields = ['status', 'checkout', 'requested_by']
    
    def get_queryset(self):
        """Filter based on user role."""
        queryset = super().get_queryset()
        user = self.request.user
        
        # Property managers and admins see all requests
        if user.role in ['SUPER_ADMIN', 'PROPERTY_MANAGER', 'MAINTENANCE_SUPERVISOR']:
            return queryset
        
        # Regular users see only their own requests
        return queryset.filter(requested_by=user)
    
    @action(detail=True, methods=['post'], permission_classes=[IsPropertyManager])
    def approve(self, request, pk=None):
        """Approve an extension request."""
        extension_request = self.get_object()
        notes = request.data.get('notes', '')
        
        try:
            extension_request.approve(request.user, notes)
            
            # Notify the requester
            from apps.core.models import Notification
            Notification.objects.create(
                user=extension_request.requested_by,
                title='Extension Request Approved',
                message=f'Your extension request for {extension_request.checkout.asset.asset_id} has been approved. New return date: {extension_request.requested_return_date}',
                notification_type='SUCCESS'
            )
            
            # Log the action
            log_action(
                user=request.user,
                action='APPROVE_EXTENSION',
                model_name='CheckoutExtensionRequest',
                object_id=extension_request.id,
                details=f'Approved extension for {extension_request.checkout.asset.asset_id}',
                ip_address=get_client_ip(request)
            )
            
            return Response({
                'message': 'Extension request approved successfully',
                'extension_request': CheckoutExtensionRequestSerializer(extension_request).data
            })
        except ValueError as e:
            return Response(
                {'error': str(e)},
                status=status.HTTP_400_BAD_REQUEST
            )
    
    @action(detail=True, methods=['post'], permission_classes=[IsPropertyManager])
    def reject(self, request, pk=None):
        """Reject an extension request."""
        extension_request = self.get_object()
        notes = request.data.get('notes', '')
        
        if not notes:
            return Response(
                {'error': 'Rejection notes are required'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        try:
            extension_request.reject(request.user, notes)
            
            # Notify the requester
            from apps.core.models import Notification
            Notification.objects.create(
                user=extension_request.requested_by,
                title='Extension Request Rejected',
                message=f'Your extension request for {extension_request.checkout.asset.asset_id} has been rejected. Reason: {notes}',
                notification_type='WARNING'
            )
            
            # Log the action
            log_action(
                user=request.user,
                action='REJECT_EXTENSION',
                model_name='CheckoutExtensionRequest',
                object_id=extension_request.id,
                details=f'Rejected extension for {extension_request.checkout.asset.asset_id}',
                ip_address=get_client_ip(request)
            )
            
            return Response({
                'message': 'Extension request rejected',
                'extension_request': CheckoutExtensionRequestSerializer(extension_request).data
            })
        except ValueError as e:
            return Response(
                {'error': str(e)},
                status=status.HTTP_400_BAD_REQUEST
            )
    
    @action(detail=False, methods=['get'])
    def statistics(self, request):
        """Get extension request statistics."""
        queryset = self.get_queryset()
        
        stats = {
            'total': queryset.count(),
            'pending': queryset.filter(status='PENDING').count(),
            'approved': queryset.filter(status='APPROVED').count(),
            'rejected': queryset.filter(status='REJECTED').count(),
        }
        
        return Response(stats)
